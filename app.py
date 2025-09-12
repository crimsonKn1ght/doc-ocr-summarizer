import os
import io
import hashlib
import streamlit as st
from PIL import Image
import pytesseract
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from sklearn.feature_extraction.text import TfidfVectorizer
import numpy as np
from langchain.embeddings.base import Embeddings

# ----------------- CONFIG ----------------- #
st.set_page_config(
    page_title="DocQ&A - Smart Document Assistant",
    page_icon="🧠",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ----------------- BACKGROUND ANIMATION ----------------- #
st.markdown(
    """
<style>
.background-svg {
    position: fixed;
    top: 0;
    left: 0;
    width: 100%;
    height: 100%;
    z-index: -1;
    opacity: 0.18; /* background transparency */
}
</style>

<svg class="background-svg" viewBox="0 0 1440 900" fill="none" xmlns="http://www.w3.org/2000/svg">
""" +
"\n".join(
    f"""
    <path d="{path}" stroke="{color}" stroke-width="2.3" stroke-linecap="round">
        <animate attributeName="stroke-dasharray" values="50 800;20 800;50 800" dur="10s" repeatCount="indefinite"/>
        <animate attributeName="stroke-dashoffset" values="800;0;800" dur="10s" repeatCount="indefinite"/>
    </path>
    """ for path, color in zip(
        [
            "M720 450C720 450 742.459 440.315 755.249 425.626C768.039 410.937 778.88 418.741 789.478 401.499C800.076 384.258 817.06 389.269 826.741 380.436C836.423 371.603 851.957 364.826 863.182 356.242C874.408 347.657 877.993 342.678 898.867 333.214C919.741 323.75 923.618 319.88 934.875 310.177C946.133 300.474 960.784 300.837 970.584 287.701C980.384 274.564 993.538 273.334 1004.85 263.087C1016.15 252.84 1026.42 250.801 1038.22 242.1C1050.02 233.399 1065.19 230.418 1074.63 215.721C1084.07 201.024 1085.49 209.128 1112.65 194.884C1139.8 180.64 1132.49 178.205 1146.43 170.636C1160.37 163.066 1168.97 158.613 1181.46 147.982C1193.95 137.35 1191.16 131.382 1217.55 125.645C1243.93 119.907 1234.19 118.899 1254.53 100.846C1274.86 82.7922 1275.12 92.8914 1290.37 76.09C1305.62 59.2886 1313.91 62.1868 1323.19 56.7536C1332.48 51.3204 1347.93 42.8082 1361.95 32.1468C1375.96 21.4855 1374.06 25.168 1397.08 10.1863C1420.09 -4.79534 1421.41 -3.16992 1431.52 -15.0078",
            "M720 450C720 450 741.044 435.759 753.062 410.636C765.079 385.514 770.541 386.148 782.73 370.489C794.918 354.83 799.378 353.188 811.338 332.597C823.298 312.005 825.578 306.419 843.707 295.493C861.837 284.568 856.194 273.248 877.376 256.48C898.558 239.713 887.536 227.843 909.648 214.958C931.759 202.073 925.133 188.092 941.063 177.621C956.994 167.151 952.171 154.663 971.197 135.041C990.222 115.418 990.785 109.375 999.488 96.1291C1008.19 82.8827 1011.4 82.2181 1032.65 61.8861C1053.9 41.5541 1045.74 48.0281 1064.01 19.5798C1082.29 -8.86844 1077.21 -3.89415 1093.7 -19.66C1110.18 -35.4258 1105.91 -46.1146 1127.68 -60.2834C1149.46 -74.4523 1144.37 -72.1024 1154.18 -97.6802C1163.99 -123.258 1165.6 -111.332 1186.21 -135.809C1206.81 -160.285 1203.29 -160.861 1220.31 -177.633C1237.33 -194.406 1236.97 -204.408 1250.42 -214.196",
            "M720 450C720 450 712.336 437.768 690.248 407.156C668.161 376.544 672.543 394.253 665.951 365.784C659.358 337.316 647.903 347.461 636.929 323.197C625.956 298.933 626.831 303.639 609.939 281.01C593.048 258.381 598.7 255.282 582.342 242.504C565.985 229.726 566.053 217.66 559.169 197.116C552.284 176.572 549.348 171.846 529.347 156.529C509.345 141.211 522.053 134.054 505.192 115.653C488.33 97.2527 482.671 82.5627 473.599 70.7833C464.527 59.0039 464.784 50.2169 447 32.0721C429.215 13.9272 436.29 0.858563 423.534 -12.6868C410.777 -26.2322 407.424 -44.0808 394.364 -56.4916C381.303 -68.9024 373.709 -72.6804 365.591 -96.1992C357.473 -119.718 358.364 -111.509 338.222 -136.495C318.08 -161.481 322.797 -149.499 315.32 -181.761C307.843 -214.023 294.563 -202.561 285.795 -223.25C277.026 -243.94 275.199 -244.055 258.602 -263.871",
            # ... (include all remaining paths from your React code, same as earlier conversion)
        ],
        [
            "#46A5CA","#8C2F2F","#4FAE4D","#D6590C","#811010","#247AFB",
            "#A534A0","#A8A438","#D6590C","#46A29C","#670F6D","#D7C200",
            "#59BBEB","#504F1C","#55BC54","#4D3568","#9F39A5","#363636",
            "#860909","#6A286F","#604483"
        ]
    )
) + """
</svg>
""",
    unsafe_allow_html=True
)

# ----------------- CUSTOM CSS ----------------- #
st.markdown(
    """
<style>
    /* Header */
    .main-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    /* (keep rest of your CSS for cards, chat, sidebar, etc.) */
</style>
""",
    unsafe_allow_html=True
)

# ----------------- UTILS ----------------- #
def get_file_hash(file_content: bytes) -> str:
    return hashlib.md5(file_content).hexdigest()

def normalize_text(text: str) -> str:
    return " ".join(text.split())

def ocr_image(image_bytes: bytes) -> str:
    try:
        image = Image.open(io.BytesIO(image_bytes))
        return pytesseract.image_to_string(image)
    except Exception:
        return ""

def extract_text_from_pdf(file_bytes: io.BytesIO, use_ocr: bool = True) -> str:
    import fitz
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        page_text = page.get_text()
        if page_text:
            text += page_text + "\n"
        if use_ocr:
            for img_meta in page.get_images(full=True):
                try:
                    base_image = doc.extract_image(img_meta[0])
                    ocr_text = ocr_image(base_image["image"])
                    if ocr_text.strip() and ocr_text not in text:
                        text += "\n" + ocr_text
                except Exception:
                    continue
    return normalize_text(text)

def extract_text_from_docx(file_bytes: io.BytesIO, use_ocr: bool = True) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_bytes)
    text = "\n".join(para.text for para in doc.paragraphs)
    if use_ocr:
        try:
            for rel in doc.part.rels.values():
                try:
                    if "image" in rel.target_ref:
                        ocr_text = ocr_image(rel.target_part.blob)
                        if ocr_text.strip() and ocr_text not in text:
                            text += "\n" + ocr_text
                except Exception:
                    continue
        except Exception:
            pass
    return normalize_text(text)

def extract_text_from_txt(file_bytes: io.BytesIO) -> str:
    try:
        text = file_bytes.read().decode("utf-8")
    except Exception:
        text = file_bytes.read().decode("utf-8", errors="ignore")
    return normalize_text(text)

# ----------------- TFIDF EMBEDDINGS ----------------- #
class TFIDFEmbeddings(Embeddings):
    def __init__(self, max_features: int = 384):
        self.vectorizer = TfidfVectorizer(
            max_features=max_features,
            stop_words="english",
            ngram_range=(1, 2),
            lowercase=True,
            token_pattern=r"\b[a-zA-Z]{2,}\b",
        )
        self.is_fitted = False
        self.dimension = max_features

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        if not self.is_fitted:
            self.vectorizer.fit(texts)
            self.is_fitted = True
        vectors = self.vectorizer.transform(texts).toarray()
        return [self._pad(v) for v in vectors]

    def embed_query(self, text: str) -> List[float]:
        if not self.is_fitted:
            return [0.0] * self.dimension
        vec = self.vectorizer.transform([text]).toarray()[0]
        return self._pad(vec)

    def _pad(self, vector: np.ndarray) -> List[float]:
        if len(vector) < self.dimension:
            vector = np.pad(vector, (0, self.dimension - len(vector)), "constant")
        return vector[: self.dimension].tolist()

# ----------------- DOCUMENT MANAGER ----------------- #
class DocumentManager:
    def __init__(self):
        self.documents: List[Document] = []
        self.processed_files = {}
        self.embeddings = TFIDFEmbeddings()
        self.vectordb = None
        try:
            self.llm = ChatGroq(model_name="llama-3.3-70b-versatile", temperature=0)
        except Exception:
            self.llm = None
        self.prompt_template = PromptTemplate(
            input_variables=["context", "question"],
            template="""Use the following context to answer the question comprehensively. If you cannot find the answer in the context, say "I cannot find the answer in the provided documents."

Context:
{context}

Question: {question}

Answer:""",
        )

    def add_file(self, filename: str, content: str, file_hash: str, file_size: int):
        if file_hash in self.processed_files:
            return False, f"File '{filename}' already processed (duplicate content)"
        if not content.strip():
            return False, f"File '{filename}' appears to be empty or unreadable"
        doc = Document(page_content=content, metadata={"source": filename, "file_hash": file_hash, "file_size": file_size})
        self.documents.append(doc)
        self.processed_files[file_hash] = {"name": filename, "size": file_size, "word_count": len(content.split())}
        return True, f"✅ Successfully processed '{filename}' ({len(content.split())} words)"

    def _rebuild_vectordb(self):
        if not self.documents:
            self.vectordb = None
            return
        try:
            splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            chunks = splitter.split_documents(self.documents)
            all_texts = [c.page_content for c in chunks]
            self.embeddings = TFIDFEmbeddings()
            _ = self.embeddings.embed_documents(all_texts)
            self.vectordb = FAISS.from_documents(chunks, self.embeddings)
        except Exception as e:
            st.error(f"Failed to build vector database: {e}")
            self.vectordb = None

    def answer_question(self, question: str) -> str:
        if not self.documents:
            return "❌ No documents uploaded. Please upload some documents first to ask questions."
        if not self.vectordb:
            return "⚠️ Document search index is not ready. Please try uploading documents again."
        if not self.llm:
            return "❌ Language model is not available. Please check your API configuration."
        try:
            docs = self.vectordb.similarity_search(question, k=5)
            if not docs:
                return "🔍 I cannot find any relevant information in the uploaded documents for your question."
            context = "\n\n".join(
                [f"📄 Source: {d.metadata.get('source','Unknown')}\n{d.page_content}" for d in docs]
            )
            chain = LLMChain(llm=self.llm, prompt=self.prompt_template)
            return chain.run(context=context, question=question)
        except Exception as e:
            return f"❌ Error processing your question: {str(e)}"

    def get_stats(self):
        total_files = len(self.processed_files)
        total_words = sum([info.get("word_count", 0) for info in self.processed_files.values()])
        total_size = sum([info.get("size", 0) for info in self.processed_files.values()])
        return {"files": total_files, "words": total_words, "size_mb": round(total_size / (1024 * 1024), 2)}

# ----------------- SESSION STATE ----------------- #
if "doc_manager" not in st.session_state:
    st.session_state.doc_manager = DocumentManager()
if "messages" not in st.session_state:
    st.session_state.messages = []

# ----------------- MAIN APP ----------------- #
st.markdown(
    """
<div class="main-header">
    <h1>🧠 DocQ&A - Smart Document Assistant</h1>
    <p style="font-size: 1.2em; margin-top: 1rem; opacity: 0.9;">
        Upload your documents and ask intelligent questions powered by AI
    </p>
</div>
""",
    unsafe_allow_html=True,
)

# Sidebar for document upload
with st.sidebar:
    st.markdown("### 📁 Document Upload")
    uploaded_files = st.file_uploader(
        "Choose your documents", type=["pdf", "docx", "txt"], accept_multiple_files=True, help="Supported formats: PDF, DOCX, TXT"
    )
    use_ocr = st.checkbox("🔍 Enable OCR for images", value=True, help="Extract text from images in documents")
    if st.button("🗑️ Clear All Documents"):
        st.session_state.doc_manager = DocumentManager()
        st.session_state.messages = []
        st.success("All documents cleared!")
        st.rerun()
    stats = st.session_state.doc_manager.get_stats()
    if stats["files"] > 0:
        st.markdown("### 📊 Document Statistics")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("📄 Files", stats["files"])
            st.metric("💾 Size (MB)", stats["size_mb"])
        with col2:
            st.metric("📝 Words", f"{stats['words']:,}")
        st.markdown("### 📋 Processed Files")
        seen_names = set()
        for file_info in st.session_state.doc_manager.processed_files.values():
            if file_info["name"] not in seen_names:
                st.markdown(f"• **{file_info['name']}** ({file_info['word_count']:,} words)")
                seen_names.add(file_info["name"])

# Process uploaded files
if uploaded_files:
    progress_bar = st.progress(0)
    status_text = st.empty()
    for i, uploaded_file in enumerate(uploaded_files):
        progress_bar.progress((i + 1) / len(uploaded_files))
        status_text.text(f"Processing {uploaded_file.name}...")
        file_data = uploaded_file.getvalue()
        file_bytes = io.BytesIO(file_data)
        file_bytes.seek(0)
        file_hash = get_file_hash(file_data)
        if file_hash in st.session_state.doc_manager.processed_files:
            continue
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        text_content = ""
        try:
            if file_extension == ".pdf":
                text_content = extract_text_from_pdf(file_bytes, use_ocr)
            elif file_extension == ".docx":
                text_content = extract_text_from_docx(file_bytes, use_ocr)
            elif file_extension == ".txt":
                text_content = extract_text_from_txt(file_bytes)
            success, message = st.session_state.doc_manager.add_file(uploaded_file.name, text_content, file_hash, uploaded_file.size)
            if success: st.success(message)
            else: st.info(message)
        except Exception as e:
            st.error(f"❌ Error processing {uploaded_file.name}: {str(e)}")
    with st.spinner("🔄 Building search index..."):
        st.session_state.doc_manager._rebuild_vectordb()
    progress_bar.empty()
    status_text.empty()
    st.success("🎉 All documents processed successfully!")

# Main chat interface
if st.session_state.doc_manager.documents:
    st.markdown("### 💬 Chat with your Documents")
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    if prompt := st.chat_input("Ask me anything about your documents..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"): st.markdown(prompt)
        with st.chat_message("assistant"):
            with st.spinner("🤔 Thinking..."):
                response = st.session_state.doc_manager.answer_question(prompt)
            st.markdown(response)
        st.session_state.messages.append({"role": "assistant", "content": response})
else:
    st.markdown(
        """
    <div class="feature-card">
        <h3>🚀 Get Started</h3>
        <p>Welcome to DocQ&A! Here's how to use this intelligent document assistant:</p>
        <ol>
            <li><strong>Upload Documents:</strong> Use the sidebar to upload PDF, DOCX, or TXT files</li>
            <li><strong>Enable OCR:</strong> Check the OCR option to extract text from images in your documents</li>
            <li><strong>Ask Questions:</strong> Once uploaded, ask any questions about your documents</li>
            <li><strong>Get Smart Answers:</strong> The AI will search through your documents and provide detailed answers</li>
        </ol>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")
st.markdown(
    """
<div class="footer-content">
    <p>🚀 Built with Streamlit • Powered by AI • Made with ❤️</p>
</div>
""", unsafe_allow_html=True)
